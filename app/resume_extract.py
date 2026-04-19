"""Extract plain text from resume files: PDF, DOCX, TXT/Markdown."""

from __future__ import annotations

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


class ResumeExtractError(Exception):
    """Raised when text cannot be extracted from the uploaded file."""


def _decode_plain_text(data: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "gbk", "gb18030"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    raise ResumeExtractError("无法将文件解码为文本，请使用 UTF-8、GBK 或上传 PDF/Word")


def extract_text_from_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ResumeExtractError("服务器未安装 PDF 解析依赖") from e
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as e:
        raise ResumeExtractError("无法读取 PDF 文件（可能已损坏）") from e
    parts: list[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text()
        except Exception:
            t = ""
        if t:
            parts.append(t)
    return "\n".join(parts)


def extract_text_from_docx(data: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError as e:
        raise ResumeExtractError("服务器未安装 Word 解析依赖") from e
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as e:
        raise ResumeExtractError("无法读取 Word 文件（请确认是否为 .docx 格式）") from e
    parts: list[str] = []
    for p in document.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _is_probably_pdf(data: bytes) -> bool:
    return len(data) >= 4 and data[:4] == b"%PDF"


def _is_probably_docx(data: bytes) -> bool:
    return len(data) >= 4 and data[:2] == b"PK" and b"word/" in data[:8000]


def extract_resume_plain_text(filename: str | None, data: bytes) -> tuple[str, list[str]]:
    """
    Return (plain_text, warnings). Warnings are user-facing hints (e.g. scanned PDF).
    """
    warnings: list[str] = []
    name = (filename or "").strip().lower()
    suffix = Path(name).suffix

    if suffix == ".doc" and not name.endswith(".docx"):
        raise ResumeExtractError(
            "暂不支持旧版 Word .doc 格式。请在 Word 中「另存为」.docx，或导出为 PDF / TXT 后上传。"
        )

    if suffix == ".pdf" or name.endswith(".pdf"):
        text = extract_text_from_pdf(data)
        if not text.strip():
            raise ResumeExtractError(
                "PDF 未提取到文字，可能是扫描版或图片型 PDF。请使用可复制文本的 PDF，或导出为 Word/TXT 后再上传。"
            )
        if len(text.strip()) < 80 and len(data) > 50_000:
            warnings.append("提取文字较短，若简历为扫描件，建议换用文本型 PDF 或 Word。")
        return text, warnings

    if suffix == ".docx" or name.endswith(".docx"):
        text = extract_text_from_docx(data)
        if not text.strip():
            raise ResumeExtractError("Word 文档中未识别到文字内容，请检查文件是否为空。")
        return text, warnings

    if suffix in (".txt", ".md", ".markdown"):
        text = _decode_plain_text(data)
        return text, warnings

    # Unknown / no extension: sniff binary formats, else plain text
    if _is_probably_pdf(data):
        text = extract_text_from_pdf(data)
        if not text.strip():
            raise ResumeExtractError(
                "PDF 未提取到文字，可能是扫描版或图片型 PDF。请换用文本型文件后上传。"
            )
        return text, warnings

    if _is_probably_docx(data):
        text = extract_text_from_docx(data)
        if not text.strip():
            raise ResumeExtractError("Word 文档中未识别到文字内容。")
        return text, warnings

    try:
        text = _decode_plain_text(data)
        return text, warnings
    except ResumeExtractError:
        raise ResumeExtractError(
            "无法识别文件格式。支持：.pdf、.docx、.txt、.md；旧版 .doc 请先另存为 .docx 或 PDF。"
        ) from None
