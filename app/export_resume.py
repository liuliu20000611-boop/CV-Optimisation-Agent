"""Build downloadable resume files (DOCX / PDF / plain text) from optimized plain text."""

from __future__ import annotations

import logging
import os
from io import BytesIO
from typing import Final

logger = logging.getLogger(__name__)

_DOCX_MAX_CHARS: Final[int] = 500_000
_PDF_MAX_CHARS: Final[int] = 500_000
# 超长单行先切段，避免 insert_textbox 自动折行占满整页高度而溢出（溢出时 PyMuPDF 整页空白）。
_PDF_MAX_LINE_CHARS: Final[int] = 400
# PyMuPDF insert_textbox 默认 fontname="helv"（Base-14 简单字体）；与 fontfile 并用时仍会按简单字体处理，
# 会将 Unicode 码点 >255 的字符替换为「?」。必须指定非 Base-14 的名称以嵌入完整 TTF/TTC。
_PDF_EMBED_FONT_NAME: Final[str] = "ResumeCJK"


def _default_cjk_font_paths() -> list[str]:
    return [
        p
        for p in (
            os.environ.get("PDF_FONT_PATH", "").strip(),
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc",
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
            r"C:\Windows\Fonts\msyh.ttc",
            r"C:\Windows\Fonts\msyhbd.ttc",
            r"C:\Windows\Fonts\simhei.ttf",
            r"C:\Windows\Fonts\simsun.ttc",
        )
        if p
    ]


def find_pdf_font() -> str | None:
    """First existing font path for CJK PDF rendering."""
    for p in _default_cjk_font_paths():
        if os.path.isfile(p):
            logger.info("PDF 使用字体: %s", p)
            return p
    return None


def _normalize_wrapped_lines(lines: list[str]) -> list[str]:
    """将超长单行切段，降低自动折行导致单页行数爆炸的概率。"""
    out: list[str] = []
    for ln in lines:
        if len(ln) <= _PDF_MAX_LINE_CHARS:
            out.append(ln)
            continue
        for i in range(0, len(ln), _PDF_MAX_LINE_CHARS):
            out.append(ln[i : i + _PDF_MAX_LINE_CHARS])
    return out


def _pdf_page_text_fits(
    text: str,
    *,
    rect: "fitz.Rect",
    font_path: str,
    fontname: str,
    fontsize: float,
) -> bool:
    """insert_textbox 在内容溢出时返回负数且不绘制任何内容，需事先探测。"""
    import fitz

    if not text:
        return True
    d = fitz.open()
    try:
        p = d.new_page(width=595, height=842)
        rc = p.insert_textbox(
            rect,
            text,
            fontname=fontname,
            fontfile=font_path,
            fontsize=fontsize,
            color=(0, 0, 0),
        )
        return rc >= 0
    finally:
        d.close()


def _split_text_for_pdf_pages(text: str, font_path: str) -> list[str]:
    """
    将正文拆成多段，每段单独一页且保证 insert_textbox 不溢出。
    换行很多的简历若按固定字数切块，仍可能整页空白（PyMuPDF 溢出则完全不绘制）。
    """
    import fitz

    rect = fitz.Rect(40, 40, 555, 802)
    fontsize = 10.5
    fn = _PDF_EMBED_FONT_NAME
    lines = _normalize_wrapped_lines(text.split("\n"))
    pages: list[str] = []
    i = 0
    while i < len(lines):
        lo, hi = i + 1, len(lines)
        best = i
        while lo <= hi:
            mid = (lo + hi + 1) // 2
            chunk = "\n".join(lines[i:mid])
            if _pdf_page_text_fits(
                chunk,
                rect=rect,
                font_path=font_path,
                fontname=fn,
                fontsize=fontsize,
            ):
                best = mid
                lo = mid + 1
            else:
                hi = mid - 1
        if best > i:
            pages.append("\n".join(lines[i:best]))
            i = best
            continue
        line = lines[i]
        lo, hi = 1, len(line)
        best = 0
        while lo <= hi:
            mid = (lo + hi + 1) // 2
            if _pdf_page_text_fits(
                line[:mid],
                rect=rect,
                font_path=font_path,
                fontname=fn,
                fontsize=fontsize,
            ):
                best = mid
                lo = mid + 1
            else:
                hi = mid - 1
        if best <= 0:
            best = 1
        pages.append(line[:best])
        if best < len(line):
            lines[i] = line[best:]
        else:
            i += 1
    return pages


def build_docx_bytes(text: str) -> bytes:
    """Create a .docx with one paragraph per line (preserve line breaks)."""
    if len(text) > _DOCX_MAX_CHARS:
        raise ValueError("正文过长，无法导出为 Word")
    import docx

    doc = docx.Document()
    if not text:
        doc.add_paragraph("")
    else:
        for line in text.split("\n"):
            doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_pdf_bytes(text: str, font_path: str) -> bytes:
    """Create a multi-page PDF; requires a CJK-capable font (TTF/TTC)."""
    if len(text) > _PDF_MAX_CHARS:
        raise ValueError("正文过长，无法导出为 PDF")
    import fitz

    doc = fitz.open()
    rect = fitz.Rect(40, 40, 555, 802)
    fontsize = 10.5
    if not text:
        doc.new_page(width=595, height=842)
        out = doc.tobytes()
        doc.close()
        return out

    for chunk in _split_text_for_pdf_pages(text, font_path):
        page = doc.new_page(width=595, height=842)
        page.insert_textbox(
            rect,
            chunk,
            fontname=_PDF_EMBED_FONT_NAME,
            fontfile=font_path,
            fontsize=fontsize,
            color=(0, 0, 0),
        )

    out = doc.tobytes()
    doc.close()
    return out


def build_plain_bytes(text: str, suffix: str) -> tuple[bytes, str]:
    raw = text.encode("utf-8")
    mime = "text/plain; charset=utf-8"
    if suffix.lower() == ".md":
        mime = "text/markdown; charset=utf-8"
    return raw, mime
